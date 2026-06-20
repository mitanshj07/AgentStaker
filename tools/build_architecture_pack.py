from __future__ import annotations

import math
from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFilter, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PACK = ROOT / "docs" / "architecture-pack"
DIAGRAMS = PACK / "diagrams"
DOCX_OUT = ROOT / "docs" / "Monad_ArenaX_Workflows_and_Architecture.docx"
HTML_OUT = PACK / "index.html"
DIAGRAMS.mkdir(parents=True, exist_ok=True)

FONT_REGULAR = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"

INK = "#292239"
MUTED = "#73677E"
CORAL = "#D95A67"
CORAL_SOFT = "#FFE8EA"
MINT = "#DDF5E5"
MINT_DARK = "#2D7B50"
PEACH = "#FFE8D5"
PEACH_DARK = "#9A5C28"
SKY = "#E2F3FF"
SKY_DARK = "#39779A"
LILAC = "#EEE3FF"
LILAC_DARK = "#7456A1"
SUN = "#FFF2C9"
SUN_DARK = "#886B12"
PANEL = "#FFFFFF"
CANVAS = "#F7F8FC"
BORDER = "#DED8E7"


def hex_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[index : index + 2], 16) for index in (0, 2, 4))


def rgba(value: str, alpha: int = 255) -> tuple[int, int, int, int]:
    return (*hex_rgb(value), alpha)


def font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def text_lines(text: str, width: int) -> list[str]:
    return wrap(text, width=max(8, width), break_long_words=False) or [""]


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, color="#A69AAF", width=4, label: str | None = None):
    draw.line([start, end], fill=rgba(color), width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 13
    wing = 0.62
    points = [
        end,
        (end[0] - size * math.cos(angle - wing), end[1] - size * math.sin(angle - wing)),
        (end[0] - size * math.cos(angle + wing), end[1] - size * math.sin(angle + wing)),
    ]
    draw.polygon(points, fill=rgba(color))
    if label:
        mid_x = int((start[0] + end[0]) / 2)
        mid_y = int((start[1] + end[1]) / 2) - 14
        label_font = font(18, True)
        box = draw.textbbox((0, 0), label, font=label_font)
        label_w = box[2] - box[0] + 18
        label_h = box[3] - box[1] + 12
        draw.rounded_rectangle((mid_x - label_w // 2, mid_y - label_h // 2, mid_x + label_w // 2, mid_y + label_h // 2), radius=10, fill=rgba("#FFFFFF", 238), outline=rgba(BORDER))
        draw.text((mid_x - label_w // 2 + 9, mid_y - label_h // 2 + 5), label, font=label_font, fill=rgba(MUTED))


def draw_lane(draw: ImageDraw.ImageDraw, box, label, fill):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=24, fill=rgba(fill, 72), outline=rgba(fill, 178), width=2)
    draw.text((x1 + 18, y1 + 12), label.upper(), font=font(18, True), fill=rgba(MUTED))


def draw_node(base: Image.Image, draw: ImageDraw.ImageDraw, box, title, detail, fill, accent=None):
    x1, y1, x2, y2 = box
    shadow = Image.new("RGBA", base.size, (0, 0, 0, 0))
    shadow_draw = ImageDraw.Draw(shadow)
    shadow_draw.rounded_rectangle((x1 + 4, y1 + 7, x2 + 4, y2 + 7), radius=18, fill=(70, 57, 89, 22))
    base.alpha_composite(shadow.filter(ImageFilter.GaussianBlur(8)))
    overlay = Image.new("RGBA", base.size, (0, 0, 0, 0))
    overlay_draw = ImageDraw.Draw(overlay)
    overlay_draw.rounded_rectangle(box, radius=18, fill=rgba(fill, 224), outline=rgba(accent or fill, 246), width=2)
    if accent:
        overlay_draw.rounded_rectangle((x1, y1, x1 + 8, y2), radius=4, fill=rgba(accent, 255))
    base.alpha_composite(overlay)
    title_font = font(21, True)
    detail_font = font(15)
    title_lines = text_lines(title, max(14, int((x2 - x1) / 13)))
    y = y1 + 14
    for line in title_lines[:2]:
        draw.text((x1 + 18, y), line, font=title_font, fill=rgba(INK))
        y += 24
    y += 4
    available_detail_lines = max(1, min(3, int((y2 - y - 11) / 18)))
    for line in text_lines(detail, max(18, int((x2 - x1) / 9)))[:available_detail_lines]:
        draw.text((x1 + 18, y), line, font=detail_font, fill=rgba(MUTED))
        y += 18


def draw_badge(draw: ImageDraw.ImageDraw, x: int, y: int, text: str, fill: str, ink: str):
    badge_font = font(17, True)
    bounds = draw.textbbox((0, 0), text, font=badge_font)
    width = bounds[2] - bounds[0] + 24
    draw.rounded_rectangle((x, y, x + width, y + 32), radius=16, fill=rgba(fill), outline=rgba(ink, 90))
    draw.text((x + 12, y + 7), text, font=badge_font, fill=rgba(ink))
    return x + width + 10


def new_canvas(title: str, subtitle: str, width=1800, height=980):
    image = Image.new("RGBA", (width, height), rgba(CANVAS))
    draw = ImageDraw.Draw(image)
    for x in range(0, width, 80):
        draw.line((x, 0, x, height), fill=rgba("#EAE6EF", 112), width=1)
    for y in range(0, height, 80):
        draw.line((0, y, width, y), fill=rgba("#EAE6EF", 112), width=1)
    draw.rounded_rectangle((34, 28, width - 34, 118), radius=22, fill=rgba("#FFFFFF", 228), outline=rgba(BORDER))
    draw.text((62, 49), title, font=font(34, True), fill=rgba(INK))
    draw.text((62, 88), subtitle, font=font(18), fill=rgba(MUTED))
    return image, draw


def save_diagram(slug: str, image: Image.Image):
    path = DIAGRAMS / f"{slug}.png"
    image.convert("RGB").save(path, quality=95)
    return path


def build_all_inclusive():
    image, draw = new_canvas("Monad ArenaX: all-inclusive architecture", "Every user surface, service, contract, event stream, and safety boundary in one view.", height=1160)
    draw_lane(draw, (42, 144, 390, 1100), "Experience", CORAL_SOFT)
    draw_lane(draw, (420, 144, 808, 1100), "Application services", SKY)
    draw_lane(draw, (840, 144, 1758, 1100), "Monad testnet protocol", MINT)
    nodes = [
        ((70, 196, 360, 304), "Sportsbook Arena", "Live pulse, categories, watchlist, quick picks, transparent ticket.", CORAL_SOFT, CORAL),
        ((70, 330, 360, 438), "Pro Exchange", "Depth, smart router, EIP-712 orders, cancellation, timeline.", PEACH, PEACH_DARK),
        ((70, 464, 360, 572), "Portfolio + Parlay NFTs", "Positions, claims, deterministic cashout, advisory hedge.", LILAC, LILAC_DARK),
        ((70, 598, 360, 706), "AI + Risk Governor", "Forecast tournament, ledgers, approval-only execution.", SKY, SKY_DARK),
        ((70, 732, 360, 840), "LP + Creator Studio", "Shared vault, rebalance proposals, creator quality checks.", MINT, MINT_DARK),
        ((70, 866, 360, 974), "Monad Cockpit", "Wallet switch, RPC latency, block pulse, explorer links.", SUN, SUN_DARK),
        ((450, 196, 778, 304), "React frontend", "Typed context store, API clients, demo fallback, responsive UI.", CORAL_SOFT, CORAL),
        ((450, 344, 778, 452), "Matcher + smart router", "REST, WebSocket, TIF validation, reservations, settlement intents.", PEACH, PEACH_DARK),
        ((450, 492, 778, 600), "Advisory AI adapter", "Model adapter when configured; deterministic explainable fallback.", LILAC, LILAC_DARK),
        ((450, 640, 778, 748), "Ponder indexer", "Contract event ingestion, indexed freshness, GraphQL-style reads.", SKY, SKY_DARK),
        ((450, 788, 778, 896), "Postgres read model", "Markets, portfolio, orderbook, agents, LP, creator revenue.", MINT, MINT_DARK),
        ((450, 936, 778, 1044), "Demo fallback fixtures", "Judge-ready continuity when RPC or optional services are slow.", SUN, SUN_DARK),
        ((876, 196, 1124, 292), "Monad RPC", "Chain ID 10143 | test MON only", SKY, SKY_DARK),
        ((1150, 196, 1428, 292), "MarketFactory.sol", "Market lifecycle + resolution state", CORAL_SOFT, CORAL),
        ((1454, 196, 1728, 292), "ResponsibleLimits.sol", "Daily, exposure, order, AI caps", SUN, SUN_DARK),
        ((876, 328, 1124, 424), "AMMPool.sol", "Outcome shares + LP accounting", MINT, MINT_DARK),
        ((1150, 328, 1428, 424), "ExchangeBook.sol", "Signature, nonce, fills, cancel", PEACH, PEACH_DARK),
        ((1454, 328, 1728, 424), "CreatorVault.sol", "Creator, referral, protocol fees", CORAL_SOFT, CORAL),
        ((876, 460, 1124, 556), "ParlayEngine.sol", "ERC-721, liability, cashout, claim", LILAC, LILAC_DARK),
        ((1150, 460, 1428, 556), "OracleCouncil.sol", "Commit, reveal, challenge, finalize", SKY, SKY_DARK),
        ((1454, 460, 1728, 556), "ForecastArena.sol", "Forecast commits + Brier scores", MINT, MINT_DARK),
        ((876, 592, 1124, 688), "SharedLiquidityVault.sol", "Deposit, queue, allocate, rebalance", PEACH, PEACH_DARK),
        ((1150, 592, 1428, 688), "AIPass.sol", "Tier pass + credit consumption", LILAC, LILAC_DARK),
        ((1454, 592, 1728, 688), "RiskGovernor.sol", "Propose, approve, execute, block", SUN, SUN_DARK),
        ((876, 724, 1124, 820), "Reputation.sol", "Authorized reporters, XP, badges", MINT, MINT_DARK),
        ((1150, 724, 1428, 820), "LeagueFactory.sol", "Community leagues + scoring", CORAL_SOFT, CORAL),
        ((1454, 724, 1728, 820), "Event stream", "Trade, fill, cashout, oracle, LP, AI pass", SKY, SKY_DARK),
        ((1008, 890, 1600, 1000), "Testnet safety envelope", "TESTNET_ONLY | POINTS_ONLY | DEMO_FALLBACK | LEGAL_REVIEW_REQUIRED. AI recommends; the wallet authorizes.", "#FFF8E2", SUN_DARK),
    ]
    edges = [
        ((360, 250), (450, 250), None), ((360, 384), (450, 398), None), ((360, 518), (450, 546), None),
        ((360, 652), (450, 546), None), ((360, 786), (450, 694), None), ((360, 920), (876, 244), "RPC"),
        ((778, 250), (876, 244), None), ((778, 398), (1150, 376), "settle"), ((778, 546), (1454, 640), "advice"),
        ((1728, 772), (778, 694), "index"), ((778, 694), (778, 842), None), ((778, 990), (450, 842), None),
        ((1124, 244), (1150, 244), None), ((1124, 376), (1150, 376), None), ((1428, 376), (1454, 376), None),
        ((1124, 508), (1150, 508), None), ((1428, 508), (1454, 508), None), ((1124, 640), (1150, 640), None),
        ((1428, 640), (1454, 640), None), ((1124, 772), (1150, 772), None), ((1428, 772), (1454, 772), None),
        ((1588, 820), (1588, 890), None), ((1508, 688), (1508, 890), None), ((1600, 292), (1600, 890), None),
    ]
    for start, end, label in edges:
        draw_arrow(draw, start, end, label=label)
    for node in nodes:
        draw_node(image, draw, *node)
    return save_diagram("01-all-inclusive-architecture", image)


def build_user_lifecycle():
    image, draw = new_canvas("End-to-end user workflow", "A judge-ready journey from launch screen to indexed settlement.", height=920)
    colors = [CORAL_SOFT, SKY, MINT, PEACH, LILAC, SUN, SKY, CORAL_SOFT, MINT, PEACH, LILAC, SUN]
    steps = [
        ("1. Launch", "Translucent Monad ArenaX fade + live pulse ticker"),
        ("2. Connect", "Add or switch Monad testnet; wallet verifies chain 10143"),
        ("3. Discover", "Trending, Live, Closing soon, New, categories, watchlist"),
        ("4. Price", "Load YES or NO quick pick into a transparent ticket"),
        ("5. Review", "Stake preset, payout preview, limits, Monad settlement label"),
        ("6. Sign", "Wallet authorizes AMM buy or EIP-712 exchange order"),
        ("7. Compose", "Add legs and mint transferable ERC-721 parlay NFT"),
        ("8. Manage", "Cashout quote, AI hedge advice, portfolio risk"),
        ("9. Resolve", "Oracle commit, reveal, challenge bond, council decision"),
        ("10. Claim", "Pull-based payout for winning shares or parlay"),
        ("11. Index", "Ponder refreshes portfolio, history, LP, leaderboard"),
        ("12. Monitor", "Monad cockpit shows block, explorer, finalized, indexed"),
    ]
    positions = []
    for row in range(3):
        for column in range(4):
            x1 = 58 + column * 434
            y1 = 188 + row * 214
            positions.append((x1, y1, x1 + 370, y1 + 126))
    for index in range(len(positions) - 1):
        current = positions[index]
        nxt = positions[index + 1]
        if (index + 1) % 4:
            draw_arrow(draw, (current[2], (current[1] + current[3]) // 2), (nxt[0], (nxt[1] + nxt[3]) // 2))
        else:
            draw_arrow(draw, ((current[0] + current[2]) // 2, current[3]), ((nxt[0] + nxt[2]) // 2, nxt[1]))
    for box, (title, detail), fill in zip(positions, steps, colors):
        draw_node(image, draw, box, title, detail, fill, CORAL if fill == CORAL_SOFT else None)
    x = 58
    x = draw_badge(draw, x, 852, "AI advisory only", LILAC, LILAC_DARK)
    x = draw_badge(draw, x, 852, "Wallet signature required", SUN, SUN_DARK)
    draw_badge(draw, x, 852, "Test MON has no monetary value", MINT, MINT_DARK)
    return save_diagram("02-user-lifecycle", image)


def build_trade_routing():
    image, draw = new_canvas("Trade execution: AMM and signed-order smart routing", "Two transparent paths converge on Monad testnet settlement and indexed confirmation.", height=950)
    draw_lane(draw, (46, 150, 364, 880), "User and UI", CORAL_SOFT)
    draw_lane(draw, (398, 150, 936, 880), "Non-custodial services", SKY)
    draw_lane(draw, (970, 150, 1750, 880), "Monad settlement", MINT)
    nodes = [
        ((76, 210, 334, 316), "Arena ticket", "Market, side, stake preset, payout preview, limits.", CORAL_SOFT, CORAL),
        ((76, 384, 334, 490), "Pro Exchange", "Limit, TIF, route preview, queue estimate, depth.", PEACH, PEACH_DARK),
        ((76, 558, 334, 664), "Wallet approval", "Transaction signature or EIP-712 order signature.", SUN, SUN_DARK),
        ((430, 230, 674, 338), "AMM quote path", "Direct deterministic AMMPool quote for outcome shares.", MINT, MINT_DARK),
        ((706, 384, 900, 492), "Smart router", "Split AMM/CLOB legs, expiry, impact, route hash.", SKY, SKY_DARK),
        ((430, 556, 674, 664), "Matcher", "Validate nonce, expiry, tick, reservations, sports lock.", PEACH, PEACH_DARK),
        ((706, 704, 900, 812), "WebSocket", "Broadcast orderbook snapshot and partial-fill updates.", LILAC, LILAC_DARK),
        ((1010, 230, 1274, 338), "ResponsibleLimits", "Exposure, spend, open-order, execution caps.", SUN, SUN_DARK),
        ((1304, 230, 1710, 338), "AMMPool.sol", "Move probabilities, emit Trade, forward creator fees.", MINT, MINT_DARK),
        ((1010, 482, 1274, 590), "ExchangeBook.sol", "Verify EIP-712, nonce, fills, cancellation, constraints.", PEACH, PEACH_DARK),
        ((1304, 482, 1710, 590), "CreatorVault.sol", "Split creator, referral, and protocol revenue.", CORAL_SOFT, CORAL),
        ((1152, 720, 1560, 828), "Ponder confirmation", "Submitted -> proposed -> finalized -> indexed.", SKY, SKY_DARK),
    ]
    edges = [
        ((334, 263), (430, 284), "quote"), ((334, 437), (706, 438), "route"), ((334, 611), (430, 610), None),
        ((552, 556), (552, 338), "AMM"), ((674, 610), (1010, 536), "intent"), ((803, 492), (552, 556), None),
        ((1274, 284), (1304, 284), None), ((1274, 536), (1304, 536), "fees"), ((1507, 338), (1507, 482), None),
        ((1507, 590), (1356, 720), "events"), ((803, 704), (552, 664), None), ((1010, 536), (1010, 284), "cap"),
    ]
    for start, end, label in edges:
        draw_arrow(draw, start, end, label=label)
    for node in nodes:
        draw_node(image, draw, *node)
    return save_diagram("03-trade-routing", image)


def build_parlay():
    image, draw = new_canvas("Parlay NFT lifecycle: reserve, cash out, hedge, claim", "Deterministic contract math remains authoritative; AI can explain or propose but never price the payout.", height=900)
    nodes = [
        ((62, 210, 336, 330), "Select 2-5 legs", "Distinct markets, YES/NO outcomes, live AMM odds.", CORAL_SOFT, CORAL),
        ((392, 210, 666, 330), "Reserve liability", "House liquidity check prevents insolvency before mint.", SUN, SUN_DARK),
        ((722, 210, 996, 330), "Mint ERC-721", "Transferable parlay position with ownerOf authorization.", LILAC, LILAC_DARK),
        ((1052, 210, 1326, 330), "Track live state", "Resolved legs + unresolved live AMM probabilities.", SKY, SKY_DARK),
        ((1382, 210, 1656, 330), "Resolve or exit", "User decides to hold, cash out, or hedge.", MINT, MINT_DARK),
        ((1046, 496, 1326, 616), "quoteCashout", "Discount bps, liquidity check, expiry, minPayout guard.", PEACH, PEACH_DARK),
        ((1382, 496, 1656, 616), "cashoutParlay", "User signature releases pull-based payout.", CORAL_SOFT, CORAL),
        ((710, 496, 990, 616), "AI hedge advice", "Explain hold-versus-exit risk and suggest separate route.", LILAC, LILAC_DARK),
        ((374, 496, 654, 616), "Signed hedge order", "User may approve a separate AMM/CLOB route.", SKY, SKY_DARK),
        ((62, 496, 318, 616), "Claim", "Winning NFT position claims after finalized resolution.", MINT, MINT_DARK),
    ]
    for start, end, label in [
        ((336, 270), (392, 270), None), ((666, 270), (722, 270), None), ((996, 270), (1052, 270), None), ((1326, 270), (1382, 270), None),
        ((1519, 330), (1519, 496), "cashout"), ((1382, 556), (1326, 556), "quote"), ((1186, 496), (850, 616), "explain"),
        ((710, 556), (654, 556), "approve"), ((374, 556), (318, 556), "resolve"),
    ]:
        draw_arrow(draw, start, end, label=label)
    for node in nodes:
        draw_node(image, draw, *node)
    draw_badge(draw, 64, 760, "AI cannot influence payout math", LILAC, LILAC_DARK)
    draw_badge(draw, 404, 760, "Cashout uses minPayout slippage guard", SUN, SUN_DARK)
    draw_badge(draw, 822, 760, "Liability reservation prevents insolvency", MINT, MINT_DARK)
    return save_diagram("04-parlay-nft", image)


def build_oracle_forecast():
    image, draw = new_canvas("Oracle Court and Forecast Arena", "Resolution integrity and agent reputation meet after the outcome is finalized.", height=940)
    draw_lane(draw, (48, 152, 872, 864), "Oracle Court", SKY)
    draw_lane(draw, (928, 152, 1752, 864), "Forecast Arena", LILAC)
    oracle_nodes = [
        ((86, 220, 354, 330), "Lock market", "Trading stops at configured deadline.", CORAL_SOFT, CORAL),
        ((462, 220, 730, 330), "Commit result", "Council publishes outcome hash.", SKY, SKY_DARK),
        ((462, 410, 730, 520), "Reveal result", "Outcome validates against market outcomes.", MINT, MINT_DARK),
        ((86, 600, 354, 710), "No challenge", "Finalize after dispute window.", MINT, MINT_DARK),
        ((462, 600, 730, 710), "Bonded challenge", "Disputed state, deadline, council decision.", PEACH, PEACH_DARK),
    ]
    forecast_nodes = [
        ((970, 220, 1238, 330), "Register agent", "Authorized participant identity and reputation.", LILAC, LILAC_DARK),
        ((1346, 220, 1614, 330), "Commit forecast", "One hashed probability per agent and market.", SKY, SKY_DARK),
        ((1346, 410, 1614, 520), "Reveal forecast", "Reveal probability before deadline.", MINT, MINT_DARK),
        ((970, 600, 1238, 710), "Brier score", "Settle accuracy after finalized outcome.", SUN, SUN_DARK),
        ((1346, 600, 1614, 710), "Leaderboard", "Badge, score, reputation report.", CORAL_SOFT, CORAL),
    ]
    edges = [
        ((354, 275), (462, 275), None), ((596, 330), (596, 410), None), ((462, 465), (354, 655), "clear"),
        ((596, 520), (596, 600), "challenge"), ((730, 655), (970, 655), "final outcome"),
        ((1238, 275), (1346, 275), None), ((1480, 330), (1480, 410), None), ((1346, 465), (1238, 655), None),
        ((1238, 655), (1346, 655), None),
    ]
    for start, end, label in edges:
        draw_arrow(draw, start, end, label=label)
    for node in oracle_nodes + forecast_nodes:
        draw_node(image, draw, *node)
    draw_badge(draw, 84, 792, "Commit-reveal", SKY, SKY_DARK)
    draw_badge(draw, 256, 792, "Challenge bond", PEACH, PEACH_DARK)
    draw_badge(draw, 974, 792, "One forecast per market", LILAC, LILAC_DARK)
    draw_badge(draw, 1254, 792, "Auditable Brier score", SUN, SUN_DARK)
    return save_diagram("05-oracle-forecast", image)


def build_ai_safety():
    image, draw = new_canvas("Agentic AI safety boundary and pass gating", "Visible ledgers, deterministic fallback, AIPass credits, and wallet approval keep AI advisory.", height=980)
    nodes = [
        ((66, 224, 340, 344), "User request", "Parlay risk, hedge, forecast, LP rebalance, or market quality.", CORAL_SOFT, CORAL),
        ((400, 224, 674, 344), "AIPass gate", "Tier, available credits, authorized AI consumer.", LILAC, LILAC_DARK),
        ((734, 224, 1008, 344), "Agent adapter", "Configured model or deterministic explainable fallback.", SKY, SKY_DARK),
        ((1068, 224, 1342, 344), "Task ledger", "Input, tool, confidence, recommendation, warning.", MINT, MINT_DARK),
        ((1402, 224, 1676, 344), "Advice", "Human-readable proposal shown in UI.", PEACH, PEACH_DARK),
        ((1068, 516, 1342, 636), "No state change", "Show analysis, alert, quality score, simulation.", MINT, MINT_DARK),
        ((1402, 516, 1676, 636), "State-changing action", "Route to explicit wallet signature.", SUN, SUN_DARK),
        ((1068, 700, 1342, 820), "ResponsibleLimits", "AI-agent execution cap can block unsafe proposals.", CORAL_SOFT, CORAL),
        ((734, 700, 1008, 820), "Monad testnet", "Only user-approved transaction may execute.", SKY, SKY_DARK),
    ]
    for start, end, label in [
        ((340, 284), (400, 284), None), ((674, 284), (734, 284), None), ((1008, 284), (1068, 284), None), ((1342, 284), (1402, 284), None),
        ((1539, 344), (1205, 516), "display"), ((1539, 344), (1539, 516), "mutate?"), ((1539, 636), (1205, 700), "cap"), ((1068, 760), (1008, 760), "sign"),
    ]:
        draw_arrow(draw, start, end, label=label)
    for node in nodes:
        draw_node(image, draw, *node)
    x = 66
    for text, fill, ink in [
        ("Free: explanation + quality", MINT, MINT_DARK), ("Pro: forecasts + risk", SKY, SKY_DARK),
        ("Creator: studio + analytics", CORAL_SOFT, CORAL), ("LP: simulations + rebalance", PEACH, PEACH_DARK),
        ("Institutional: API + batch", LILAC, LILAC_DARK),
    ]:
        x = draw_badge(draw, x, 894, text, fill, ink)
    return save_diagram("06-ai-safety", image)


def build_lp_creator():
    image, draw = new_canvas("Shared LP vault and creator economy", "Liquidity deployment, queued withdrawals, AI rebalancing, and fee distribution stay visible.", height=920)
    draw_lane(draw, (46, 154, 938, 858), "Shared liquidity", MINT)
    draw_lane(draw, (968, 154, 1752, 858), "Creator economy", CORAL_SOFT)
    nodes = [
        ((86, 220, 336, 330), "LP deposits test MON", "Mint proportional vault shares.", MINT, MINT_DARK),
        ((402, 220, 652, 330), "Idle liquidity", "Available for withdrawal or allocation.", SKY, SKY_DARK),
        ((718, 220, 898, 330), "Operator allocation", "Deploy capped liquidity to AMM markets.", PEACH, PEACH_DARK),
        ((718, 452, 898, 562), "AI rebalance proposal", "Drawdown simulation and suggested move.", LILAC, LILAC_DARK),
        ((402, 452, 652, 562), "Rebalance", "Recall and deploy with minReceived guard.", SUN, SUN_DARK),
        ((86, 452, 336, 562), "Queued withdrawal", "Queue when idle liquidity is insufficient.", CORAL_SOFT, CORAL),
        ((1008, 220, 1266, 330), "AMM buy / sell", "Creator-aware and referrer-aware entrypoints.", PEACH, PEACH_DARK),
        ((1326, 220, 1690, 330), "CreatorVault.sol", "Authorized fee sink for AMM and exchange.", CORAL_SOFT, CORAL),
        ((1008, 452, 1266, 562), "Creator balance", "Claimable market-creator revenue.", MINT, MINT_DARK),
        ((1326, 452, 1690, 562), "Referral balance", "Claimable referral attribution.", SKY, SKY_DARK),
        ((1168, 674, 1532, 784), "Protocol balance", "Transparent fee split supports operations.", LILAC, LILAC_DARK),
    ]
    for start, end, label in [
        ((336, 275), (402, 275), None), ((652, 275), (718, 275), None), ((808, 330), (808, 452), "monitor"),
        ((718, 507), (652, 507), "approve"), ((402, 507), (336, 507), "recall"), ((211, 330), (211, 452), "insufficient idle"),
        ((898, 275), (1008, 275), "depth"), ((1266, 275), (1326, 275), "fees"), ((1508, 330), (1137, 452), "creator"),
        ((1508, 330), (1508, 452), "referral"), ((1508, 562), (1350, 674), "protocol"),
    ]:
        draw_arrow(draw, start, end, label=label)
    for node in nodes:
        draw_node(image, draw, *node)
    return save_diagram("07-lp-creator", image)


def build_indexer_ops():
    image, draw = new_canvas("Indexer, runtime modes, and operational observability", "Monad events become fresh product reads while the app retains graceful demo continuity.", height=920)
    nodes = [
        ((64, 232, 342, 352), "Monad contracts", "Trade, fill, cancel, cashout, forecast, court, LP, pass, revenue.", MINT, MINT_DARK),
        ((404, 232, 682, 352), "Ponder handlers", "Decode event-specific records and indexed freshness.", SKY, SKY_DARK),
        ((744, 232, 1022, 352), "Postgres", "Queryable protocol read model.", LILAC, LILAC_DARK),
        ((1084, 232, 1362, 352), "Indexed API", "Markets, portfolio, leaderboard, LP, orders, agents.", PEACH, PEACH_DARK),
        ((1424, 232, 1702, 352), "React services", "Typed clients with local fixture fallback.", CORAL_SOFT, CORAL),
        ((1424, 532, 1702, 652), "Monad cockpit", "Block, RPC latency, wallet health, explorer, timeline.", SKY, SKY_DARK),
        ((1084, 532, 1362, 652), "Demo fallback", "Deterministic fixtures preserve judge flow.", SUN, SUN_DARK),
        ((744, 532, 1022, 652), "Safety modes", "TESTNET_ONLY, POINTS_ONLY, LEGAL_REVIEW_REQUIRED.", CORAL_SOFT, CORAL),
        ((404, 532, 682, 652), "Matcher health", "REST, WebSocket, router, signed cancel, sports lock.", PEACH, PEACH_DARK),
        ((64, 532, 342, 652), "Verification", "Lint, build, codegen, Foundry tests, matcher smoke.", MINT, MINT_DARK),
    ]
    for start, end, label in [
        ((342, 292), (404, 292), "events"), ((682, 292), (744, 292), None), ((1022, 292), (1084, 292), None), ((1362, 292), (1424, 292), None),
        ((1563, 352), (1563, 532), "status"), ((1424, 592), (1362, 592), "fallback"), ((1084, 592), (1022, 592), None),
        ((744, 592), (682, 592), "enforce"), ((404, 592), (342, 592), "smoke"),
    ]:
        draw_arrow(draw, start, end, label=label)
    for node in nodes:
        draw_node(image, draw, *node)
    return save_diagram("08-indexer-operations", image)


def build_diagrams():
    return [
        build_all_inclusive(),
        build_user_lifecycle(),
        build_trade_routing(),
        build_parlay(),
        build_oracle_forecast(),
        build_ai_safety(),
        build_lp_creator(),
        build_indexer_ops(),
    ]


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_table_header(table):
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_run(paragraph, text, *, size=10.5, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*hex_rgb(color))
    run.bold = bold
    run.italic = italic
    return run


def add_text(doc, text, *, size=10.5, color=INK, bold=False, italic=False, after=5, before=0, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.18
    if align is not None:
        paragraph.alignment = align
    add_run(paragraph, text, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_bullet(doc, text, after=3):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.18
    add_run(paragraph, text, size=10.2, color=INK)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    add_run(paragraph, text, size={1: 16, 2: 13, 3: 11.5}[level], color={1: CORAL, 2: LILAC_DARK, 3: SKY_DARK}[level], bold=True)
    return paragraph


def add_callout(doc, title: str, text: str, fill=MINT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    mark_table_header(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title.upper(), size=9, color=MINT_DARK if fill == MINT else LILAC_DARK, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, text, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    mark_table_header(table)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "#EEE8F8")
        p = cell.paragraphs[0]
        add_run(p, header, size=9.3, color=INK, bold=True)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if row_index % 2:
                set_cell_shading(cells[index], "#FBFAFD")
            p = cells[index].paragraphs[0]
            add_run(p, value, size=9.1, color=INK, bold=index == 0)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_picture_with_alt(document, path: Path, alt: str):
    picture = document.add_picture(str(path), width=Inches(6.48))
    picture._inline.docPr.set("descr", alt)
    picture._inline.docPr.set("title", alt)
    return picture


def build_docx(diagrams: list[Path]):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)
    document.core_properties.title = "Monad ArenaX Workflows and Architecture"
    document.core_properties.subject = "Monad testnet prediction exchange architecture and workflow design pack"
    document.core_properties.author = "Monad ArenaX"
    document.core_properties.keywords = "Monad, prediction market, architecture, workflow, hackathon, testnet"

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(*hex_rgb(INK))
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18
    for level, size, color, before, after in [
        (1, 16, CORAL, 16, 8),
        (2, 13, LILAC_DARK, 12, 6),
        (3, 11.5, SKY_DARK, 9, 4),
    ]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*hex_rgb(color))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    add_run(header, "MONAD ARENAX  /  WORKFLOW & ARCHITECTURE PACK", size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_run(footer, "Monad testnet only  |  test MON has no monetary value  |  ", size=8.2, color=MUTED)
    add_page_number(footer)

    add_text(document, "MONAD ARENAX", size=11, color=CORAL, bold=True, after=56)
    add_text(document, "Workflows &\nArchitecture", size=30, color=INK, bold=True, after=9)
    add_text(document, "A complete visual design pack for the Monad-first prediction exchange", size=14, color=LILAC_DARK, bold=True, after=14)
    add_text(document, "Hackathon testnet edition  |  Chain ID 10143  |  May 31, 2026", size=10.5, color=MUTED, after=22)
    add_callout(document, "Design intent", "Pastel translucent diagrams, operationally precise workflows, and one all-inclusive system map. Every value-moving path remains testnet-only and every AI action remains advisory until the user signs.", LILAC)
    add_matrix(document, ["DOCUMENT MAP", "COVERAGE"], [
        ["System map", "Frontend, services, Monad contracts, events, safety modes, fallback"],
        ["Product workflows", "Discovery, trading, Pro Exchange, parlay NFTs, oracle, AI, LP, creator"],
        ["Operations", "Indexer reads, matcher health, runtime modes, verification"],
    ], [2200, 7160])
    add_text(document, "Prepared for hackathon judges, technical reviewers, and product walkthroughs.", size=10.5, color=MUTED, italic=True, before=6)
    document.add_page_break()

    add_heading(document, "1. Executive Architecture Map", 1)
    add_text(document, "The map below is the one inclusive reference for the complete ArenaX system. Read left to right: user-facing product surfaces, supporting application services, Monad testnet protocol modules, indexed reads, and the safety envelope.")
    add_picture_with_alt(document, diagrams[0], "All-inclusive Monad ArenaX architecture showing experience surfaces, application services, Monad testnet contracts, indexed events, and safety modes.")
    add_text(document, "Figure 1. All-inclusive Monad ArenaX architecture.", size=8.6, color=MUTED, italic=True, after=5)
    add_callout(document, "Non-negotiable boundary", "The matcher never custodies user funds. AI never places an order. Deterministic contract math and explicit wallet approval remain authoritative.", MINT)
    document.add_page_break()

    sections = [
        ("2. End-to-End Product Journey", "The product experience is structured as one continuous testnet journey. It begins with the branded Arena launch, then moves through discovery, review-first pricing, wallet approval, portfolio management, oracle resolution, and indexed confirmation.", diagrams[1], [
            "The Arena view opens first so judges see the usable prediction exchange immediately.",
            "Quick picks prepare a transparent ticket; they do not submit a transaction.",
            "The finality timeline distinguishes submitted, proposed, finalized, and indexed states.",
        ]),
        ("3. AMM and Pro Exchange Execution", "ArenaX supports a simple AMM path and a professional signed-order path. Both are non-custodial, visible to the user, and settled on Monad testnet.", diagrams[2], [
            "AMM trades move outcome-share probabilities and forward creator-aware fees.",
            "The smart router can preview AMM/CLOB splits, impact, expiry, and route hash.",
            "The matcher validates time-in-force, nonce, tick size, reservations, and sports locks before it creates a settlement intent.",
        ]),
        ("4. Parlay NFT Cashout and Hedge", "Parlays are transferable ERC-721 positions. The contract reserves liability before minting and computes early exits from deterministic AMM probabilities.", diagrams[3], [
            "Cashout math applies unresolved probabilities, resolved-leg state, discount basis points, available liquidity, and a minPayout guard.",
            "The AI layer can explain hold-versus-cashout risk and propose a hedge route, but cannot modify payout math.",
            "Claims remain pull-based after finalized resolution.",
        ]),
        ("5. Oracle Court and Forecast Arena", "Outcome integrity is visible: Oracle Court handles commit-reveal and bonded disputes, while Forecast Arena turns finalized outcomes into auditable Brier scores and reputation updates.", diagrams[4], [
            "Challenges carry a configured bond and place the market into a disputed state.",
            "Forecast agents reveal one probability per market before deadline.",
            "Finalized results drive leaderboards, badges, and authorized reputation reports.",
        ]),
        ("6. Agentic AI Safety and Tier Gating", "ArenaX exposes AI as a visible advisory layer rather than an invisible executor. Every task writes a ledger and every state-changing proposal routes back to the user.", diagrams[5], [
            "AIPass tiers gate credits and feature access for Free, Pro, Creator, LP, and Institutional users.",
            "When a model key is unavailable, deterministic fallback analysis preserves an explainable demo.",
            "ResponsibleLimits can block unsafe AI execution proposals before wallet approval.",
        ]),
        ("7. Shared LP Vault and Creator Economy", "The liquidity workflow is designed for transparent deployment and safe recall. The creator economy is wired into trading rather than shown as a disconnected dashboard.", diagrams[6], [
            "Vault shares represent proportional deposits; withdrawals queue when idle liquidity is insufficient.",
            "The LP agent proposes rebalances with drawdown context; the operator still approves execution.",
            "CreatorVault separates creator, referral, and protocol balances for pull-based claims.",
        ]),
        ("8. Indexer, Runtime Modes, and Operations", "Ponder converts Monad events into a queryable product read model. Typed frontend services read indexed APIs and retain local fallback fixtures for hackathon continuity.", diagrams[7], [
            "Indexed domains include markets, exchange, parlays, forecasts, oracle cases, LP vault activity, passes, leagues, badges, spend, creator revenue, and risk proposals.",
            "Runtime modes remain explicit: TESTNET_ONLY, POINTS_ONLY, DEMO_FALLBACK, and LEGAL_REVIEW_REQUIRED.",
            "The verification command covers frontend lint/build, matcher TypeScript build, Ponder codegen, and Foundry tests.",
        ]),
    ]
    for title, intro, diagram, bullets in sections:
        add_heading(document, title, 1)
        add_text(document, intro)
        add_picture_with_alt(document, diagram, title.split(". ", 1)[1])
        add_text(document, f"Figure {sections.index((title, intro, diagram, bullets)) + 2}. {title.split('. ', 1)[1]}.", size=8.6, color=MUTED, italic=True, after=4)
        for bullet in bullets:
            add_bullet(document, bullet)
        document.add_page_break()

    add_heading(document, "9. Contract and Event Completeness Matrix", 1)
    add_text(document, "This matrix confirms that the visual workflows cover each implemented protocol module and its most important emitted or indexed behavior.")
    add_matrix(document, ["MODULE", "PRIMARY RESPONSIBILITY", "INDEXED OR VISIBLE BEHAVIOR"], [
        ["MarketFactory", "Market registry and lifecycle", "MarketCreated, locked/disputed/finalized state"],
        ["AMMPool", "Outcome shares and LP accounting", "Trade, quote movement, creator-aware fees"],
        ["ExchangeBook", "EIP-712 settlement", "Fill, cancellation, nonce and partial-fill state"],
        ["ParlayEngine", "ERC-721 parlay positions", "Mint, liability reserve, cashout, claim"],
        ["OracleCouncil", "Resolution integrity", "Commit, reveal, challenge, council finalize"],
        ["ForecastArena", "Agent tournament", "Agent registration, forecast commit/reveal, Brier score"],
        ["SharedLiquidityVault", "Shared LP capital", "Deposit, allocation, queued withdrawal, rebalance"],
        ["CreatorVault", "Revenue accounting", "Creator, referral, and protocol claimable balances"],
        ["AIPass", "Paid AI access", "Tier mint, credit spend, authorized consumer"],
        ["ResponsibleLimits", "Usage controls", "Exposure, daily usage, open order, execution caps"],
        ["RiskGovernor", "Approval-only risk automation", "Proposal, approval, execution, block"],
        ["Reputation + LeagueFactory", "Community layer", "League join, score, XP, badge issuance"],
    ], [1700, 3300, 4360])
    add_heading(document, "Indexed Event Coverage", 2)
    add_matrix(document, ["DOMAIN", "EVENT COVERAGE"], [
        ["Markets", "Market created, trade, result finalized"],
        ["Exchange", "Fill, cancellation, routed settlement"],
        ["Parlays", "NFT creation, liability reservation, cashout, claim"],
        ["Forecast Arena", "Registration, commit, reveal, score"],
        ["Oracle Court", "Commit, reveal, challenge, resolution"],
        ["Shared LP vault", "Deposit, allocation, queued withdrawal, rebalance"],
        ["AI passes", "Tier mint, spend record, credit consumption"],
        ["Community", "League join, badge issued, reputation report"],
        ["Creator economy", "Creator fee, referral credit, protocol balance"],
    ], [2100, 7260])
    document.add_page_break()

    add_heading(document, "10. Judge Walkthrough and Verification", 1)
    add_text(document, "The strongest demo is short, linear, and visibly Monad-native. Use the built-in runbook as a backup, while presenting Arena as the primary product experience.")
    runbook = [
        "Open ArenaX and let the translucent launch sequence fade into the live market board.",
        "Connect or switch the wallet to Monad testnet and point out chain ID 10143, RPC health, block pulse, and explorer links.",
        "Filter Live markets, load a YES or NO quick pick, review payout and responsible-limit status, then sign the testnet trade.",
        "Add two legs and mint the transferable parlay NFT. Open Portfolio and request deterministic cashout plus advisory hedge analysis.",
        "Open Pro Exchange, preview AMM/CLOB routing, create a signed-order preview, and show the settlement timeline.",
        "Run Forecast Arena, advance an Oracle Court challenge, approve an LP vault rebalance, and create an AI-reviewed market draft.",
        "Close on indexed freshness, the safety modes, and the fact that test MON has no monetary value.",
    ]
    for index, item in enumerate(runbook, 1):
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        add_run(paragraph, item, size=10.2, color=INK)
    add_callout(document, "One-command verification", "Run npm run verify for frontend lint/build, matcher TypeScript build, Ponder codegen, and all Foundry tests. Run npm run matcher:smoke while the matcher is live for router, advisory AI, order types, cancellation, sports lock, WebSocket, nonce, and rate-limit checks.", SKY)
    add_heading(document, "Runtime Safety Modes", 2)
    add_matrix(document, ["MODE", "PURPOSE"], [
        ["TESTNET_ONLY", "Default hackathon mode. Transactions use Monad testnet only."],
        ["POINTS_ONLY", "Disables wallet value flows while preserving the product walkthrough."],
        ["DEMO_FALLBACK", "Uses deterministic fixtures when optional APIs are unavailable."],
        ["LEGAL_REVIEW_REQUIRED", "Keeps any future real-money path disabled pending review."],
    ], [2200, 7160])
    add_text(document, "End of visual architecture pack.", size=9, color=MUTED, italic=True, before=8)

    document.save(DOCX_OUT)


def build_html(diagrams: list[Path]):
    cards = [
        ("01", "All-inclusive system map", "Experience, services, Monad contracts, events, safety, and fallback.", diagrams[0].name),
        ("02", "End-to-end product journey", "From branded launch to indexed settlement.", diagrams[1].name),
        ("03", "AMM and Pro Exchange routing", "Transparent execution and non-custodial settlement.", diagrams[2].name),
        ("04", "Parlay NFT lifecycle", "Liability reserve, cashout, hedge, and claim.", diagrams[3].name),
        ("05", "Oracle Court + Forecast Arena", "Bonded dispute resolution and Brier-score reputation.", diagrams[4].name),
        ("06", "Agentic AI safety boundary", "AIPass gating, visible ledgers, and wallet approval.", diagrams[5].name),
        ("07", "LP vault + creator economy", "Shared liquidity, rebalancing, and transparent fee splits.", diagrams[6].name),
        ("08", "Indexer + operations", "Ponder reads, runtime modes, cockpit, and verification.", diagrams[7].name),
    ]
    navigation = "".join(f'<a href="#flow-{number}">{number}</a>' for number, *_ in cards)
    figures = "".join(
        f"""
        <article class="flow-card" id="flow-{number}">
          <div class="flow-copy">
            <span>{number}</span>
            <div><h2>{title}</h2><p>{description}</p></div>
          </div>
          <img src="diagrams/{filename}" alt="{title}">
        </article>
        """
        for number, title, description, filename in cards
    )
    html = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Monad ArenaX | Workflow & Architecture Board</title>
  <style>
    :root {{ --ink:#292239; --muted:#73677e; --coral:#d95a67; --line:#ded8e7; --canvas:#f7f8fc; }}
    * {{ box-sizing:border-box; }}
    html {{ scroll-behavior:smooth; }}
    body {{ margin:0; background:#f7f8fc; color:var(--ink); font-family:Arial, sans-serif; }}
    body::before {{ content:""; position:fixed; inset:0; pointer-events:none; opacity:.52; background-size:48px 48px; background-image:linear-gradient(#ebe7ef 1px,transparent 1px),linear-gradient(90deg,#ebe7ef 1px,transparent 1px); }}
    .shell {{ position:relative; width:min(1480px,calc(100% - 28px)); margin:0 auto; padding:20px 0 56px; }}
    .masthead,.summary,.flow-card,.matrix {{ border:1px solid rgba(115,103,126,.18); border-radius:14px; background:rgba(255,255,255,.82); box-shadow:0 16px 44px rgba(78,63,95,.08); backdrop-filter:blur(16px); }}
    .masthead {{ padding:28px; }}
    .eyebrow {{ color:var(--coral); font-size:12px; font-weight:800; text-transform:uppercase; }}
    h1 {{ max-width:960px; margin:8px 0 6px; font-size:clamp(44px,7vw,92px); line-height:.96; }}
    h1 b {{ color:var(--coral); }}
    p {{ color:var(--muted); line-height:1.48; }}
    .meta {{ display:flex; gap:8px; flex-wrap:wrap; margin-top:18px; }}
    .meta span {{ padding:7px 10px; border-radius:999px; font-size:12px; font-weight:800; }}
    .mint {{ background:#ddf5e5; color:#2d7b50; }} .sky {{ background:#e2f3ff; color:#39779a; }} .lilac {{ background:#eee3ff; color:#7456a1; }} .peach {{ background:#ffe8d5; color:#9a5c28; }}
    .summary {{ display:grid; grid-template-columns:repeat(4,1fr); gap:1px; margin-top:14px; overflow:hidden; }}
    .summary div {{ min-height:98px; padding:18px; background:rgba(255,255,255,.72); }}
    .summary b {{ display:block; margin-top:8px; font-size:24px; }}
    nav {{ position:sticky; z-index:3; top:10px; display:flex; gap:8px; width:max-content; max-width:100%; margin:16px auto; padding:7px; overflow-x:auto; border:1px solid rgba(115,103,126,.16); border-radius:999px; background:rgba(255,255,255,.82); backdrop-filter:blur(14px); }}
    nav a {{ display:grid; place-items:center; width:34px; height:34px; border-radius:50%; color:#6f607a; font-size:12px; font-weight:800; text-decoration:none; }}
    nav a:hover {{ background:#ffe8ea; color:#c94956; }}
    .flow-card {{ margin-top:16px; padding:16px; scroll-margin-top:74px; }}
    .flow-copy {{ display:flex; align-items:center; gap:12px; padding:2px 4px 14px; }}
    .flow-copy>span {{ display:grid; place-items:center; width:38px; height:38px; border-radius:11px; background:#ffe8ea; color:#c94956; font-weight:800; }}
    h2 {{ margin:0; font-size:22px; }} .flow-copy p {{ margin:4px 0 0; }}
    img {{ display:block; width:100%; border:1px solid rgba(115,103,126,.14); border-radius:10px; }}
    .matrix {{ margin-top:16px; padding:20px; }}
    .matrix-grid {{ display:grid; grid-template-columns:repeat(4,1fr); gap:8px; }}
    .matrix-grid div {{ padding:13px; border:1px solid rgba(115,103,126,.13); border-radius:10px; background:rgba(255,255,255,.66); }}
    .matrix-grid b {{ display:block; color:#5e4c6d; font-size:13px; }} .matrix-grid small {{ color:#786b82; }}
    footer {{ padding:24px 6px 0; color:#776a80; font-size:12px; }}
    @media(max-width:820px) {{ .summary,.matrix-grid {{ grid-template-columns:1fr 1fr; }} .masthead {{ padding:20px; }} h1 {{ font-size:50px; }} }}
    @media(max-width:520px) {{ .summary,.matrix-grid {{ grid-template-columns:1fr; }} }}
    @media print {{
      @page {{ size:landscape; margin:10mm; }}
      body {{ background:white; -webkit-print-color-adjust:exact; print-color-adjust:exact; }}
      body::before,nav {{ display:none; }}
      .shell {{ width:100%; padding:0; }}
      .masthead {{ min-height:178mm; display:grid; align-content:center; page-break-after:always; }}
      .summary,nav {{ display:none; }}
      .flow-card {{ min-height:176mm; display:grid; align-content:start; margin:0; page-break-after:always; box-shadow:none; }}
      .flow-card img {{ max-height:150mm; object-fit:contain; }}
      .matrix {{ min-height:176mm; page-break-after:auto; box-shadow:none; }}
    }}
  </style>
</head>
<body>
  <main class="shell">
    <section class="masthead">
      <span class="eyebrow">Monad ArenaX / Testnet workflow design pack</span>
      <h1>Complete workflows & <b>architecture</b></h1>
      <p>One inclusive protocol map plus focused visual workflows for the full Monad-first prediction exchange: discovery, trading, signed orders, parlay NFTs, Oracle Court, Forecast Arena, agentic AI, shared liquidity, creator revenue, indexing, runtime modes, and verification.</p>
      <div class="meta"><span class="mint">Chain ID 10143</span><span class="sky">AI advisory only</span><span class="lilac">Wallet signature required</span><span class="peach">Test MON has no monetary value</span></div>
    </section>
    <section class="summary"><div><span class="eyebrow">Protocol</span><b>13 modules</b><p>Monad testnet contracts</p></div><div><span class="eyebrow">Workflows</span><b>8 visual maps</b><p>Product and operations</p></div><div><span class="eyebrow">Data</span><b>Ponder + Postgres</b><p>Indexed read model</p></div><div><span class="eyebrow">Safety</span><b>4 runtime modes</b><p>Testnet-first guardrails</p></div></section>
    <nav aria-label="Workflow navigation">{navigation}</nav>
    {figures}
    <section class="matrix"><span class="eyebrow">Completeness matrix</span><h2>Nothing important is left off the map</h2><p>Every implemented product domain is represented in the diagrams and linked back to its Monad-native protocol surface.</p><div class="matrix-grid"><div><b>Markets</b><small>Factory, AMM, sportsbook board</small></div><div><b>Exchange</b><small>Router, matcher, signed settlement</small></div><div><b>Positions</b><small>Portfolio, parlay NFT, cashout</small></div><div><b>Integrity</b><small>Oracle Court, forecasts, reputation</small></div><div><b>AI</b><small>AIPass, task ledgers, governor</small></div><div><b>Liquidity</b><small>Shared vault, queue, rebalance</small></div><div><b>Economy</b><small>Creator, referral, protocol fees</small></div><div><b>Operations</b><small>Indexer, cockpit, modes, verify</small></div></div></section>
    <footer>Monad ArenaX / hackathon testnet edition / generated visual architecture board</footer>
  </main>
</body>
</html>"""
    HTML_OUT.write_text(html)


if __name__ == "__main__":
    paths = build_diagrams()
    build_docx(paths)
    build_html(paths)
    print(f"Built {len(paths)} diagrams")
    print(DOCX_OUT)
    print(HTML_OUT)
